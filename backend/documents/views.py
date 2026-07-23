import datetime
import io
import markdown
import re
from django.http import FileResponse
from rest_framework import viewsets, permissions, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.decorators import action
from docx import Document

# Pure-python PDF library (robust across environments)
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from .models import BusinessDocument, DocumentApprovalHistory, DocumentVersion
from .serializers import BusinessDocumentSerializer
from projects.models import Project
from core.responses import api_success, api_error
from core.exceptions import ValidationError

# AI Integration for document enhancement
try:
    from srs.ai_integration import ai_service
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False

class BusinessDocumentViewSet(viewsets.ModelViewSet):
    """
    ViewSet handling BusinessDocument CRUD operations.
    Supports automated BRD/FRD generation and official sign-off workflows.
    """
    serializer_class = BusinessDocumentSerializer
    permission_classes = [IsAuthenticated]

    def get_permissions(self):
        if self.action in ["create", "update", "partial_update", "destroy", "generate_document"]:
            from billing.permissions import IsProOrEnterprise
            return [IsAuthenticated(), IsProOrEnterprise()]
        return [IsAuthenticated()]


    def get_queryset(self):
        user = self.request.user
        if not user.is_authenticated or not user.organization_id:
            return BusinessDocument.objects.none()

        queryset = BusinessDocument.objects.filter(project__organization_id=user.organization_id)
        
        # Support ?project=uuid query filtering
        project_id = self.request.query_params.get("project")
        if project_id:
            queryset = queryset.filter(project_id=project_id)
            
        # Support ?doc_type=BRD query filtering
        doc_type = self.request.query_params.get("doc_type")
        if doc_type:
            queryset = queryset.filter(doc_type=doc_type)
            
        return queryset

    def perform_create(self, serializer):
        doc = serializer.save(created_by=self.request.user)
        # Create a document version snapshot representing the initial version state
        DocumentVersion.objects.create(
            document=doc,
            version=doc.version,
            content=doc.content,
            created_by=self.request.user
        )

    def perform_update(self, serializer):
        old_doc = self.get_object()
        old_content = old_doc.content
        old_version = old_doc.version
        
        new_doc = serializer.save()
        
        # If content or version changed, create a history snapshot version
        if new_doc.content != old_content or new_doc.version != old_version:
            DocumentVersion.objects.create(
                document=new_doc,
                version=old_version,
                content=old_content,
                created_by=self.request.user
            )

    @action(detail=False, methods=["post"], url_path="generate")
    def generate_document(self, request):
        """
        Synthesizes a markdown BRD/FRD template.
        Gathers all stakeholders, requirements, and user stories.
        """
        project_id = request.data.get("project")
        doc_type = request.data.get("doc_type", "BRD")
        if not project_id:
            return api_error(message="Project ID is required.")

        try:
            project = Project.objects.get(
                id=project_id,
                organization_id=request.user.organization_id
            )
        except Project.DoesNotExist:
            return api_error(message="Project not found in your organization.")

        stakeholders = project.stakeholders.all()
        requirements = project.requirements.all()

        if doc_type == "SWOT":
            title = f"SWOT Analysis Charter - {project.name} - Version 1.0"
            content = (
                f"# SWOT Analysis Charter: {project.name}\n\n"
                f"## 1. Executive Summary & Goals\n"
                f"- **Project**: {project.name}\n"
                f"- **Workspace Organization**: {project.organization.name}\n"
                f"- **Author**: @{request.user.username}\n"
                f"- **Status**: DRAFT\n\n"
                f"## 2. SWOT Matrix & Guided Parameters\n\n"
                f"### Strengths (Internal & Helpful)\n"
                f"- **Key advantage 1**: [Describe proprietary tech, experienced team, or strong backing]\n"
                f"- **Key advantage 2**: [Detail cost efficiency or high product quality factors]\n\n"
                f"### Weaknesses (Internal & Harmful)\n"
                f"- **Constraint 1**: [Acknowledge budget limits, resource gaps, or dependencies]\n"
                f"- **Constraint 2**: [Note tech debt or team velocity blocks]\n\n"
                f"### Opportunities (External & Helpful)\n"
                f"- **Market opportunity 1**: [Highlight target niches, gaps in competitor offerings]\n"
                f"- **Market opportunity 2**: [Specify regulatory shifts or compliance openings]\n\n"
                f"### Threats (External & Harmful)\n"
                f"- **Threat 1**: [Identify macro competition or pricing pressures]\n"
                f"- **Threat 2**: [Examine cyber security risks or supply issues]\n\n"
                f"## 3. Actionable Mitigation & Growth Strategies\n"
                f"- **S-O Strategies (Maximize Strengths to Seize Opportunities)**: ...\n"
                f"- **W-T Strategies (Minimize Weaknesses to Prevent Threats)**: ...\n"
            )
        elif doc_type == "GAP":
            title = f"GAP Analysis Framework - {project.name} - Version 1.0"
            content = (
                f"# GAP Analysis Framework Document: {project.name}\n\n"
                f"## 1. Scope & Objective\n"
                f"- **Project Scope**: {project.name}\n"
                f"- **Author**: @{request.user.username}\n"
                f"- **Status**: DRAFT\n\n"
                f"## 2. Current State (As-Is) vs. Future State (To-Be)\n\n"
                f"### Business Area 1: User Onboarding Flow\n"
                f"- **Current State (As-Is)**: [Describe current manual, slow, or friction-prone setup]\n"
                f"- **Desired Future State (To-Be)**: [Target automated, biometric login, or rapid invite flow]\n"
                f"- **Gap Identified**: [State missing APIs, UI pages, or credentials]\n"
                f"- **Action Plan**: [Assign requirements or stories to resolve the gap]\n\n"
                f"### Business Area 2: Meeting & Invitations Channels\n"
                f"- **Current State (As-Is)**: [Manual messaging, no automated invites or virtual link]\n"
                f"- **Desired Future State (To-Be)**: [Jitsi virtual meeting URLs emailed automatically to invitees]\n"
                f"- **Gap Identified**: [Need integrated email dispatch hooks]\n"
                f"- **Action Plan**: [Deploy Django mailer trigger and Jitsi templates]\n\n"
                f"## 3. Resource & Capability Mapping\n"
                f"| Gap Reference | Effort Estimate | Priority | Responsible Lead |\n"
                f"| --- | --- | --- | --- |\n"
                f"| Biometric Credentials Gap | Medium | High | John Doe |\n"
                f"| Email dispatch channels | Low | High | Sarah Connor |\n"
            )
        elif doc_type == "IEEE":
            title = f"IEEE Standard Document - {project.name} - Version 1.0"
            content = (
                f"# IEEE Standard Document: {project.name}\n\n"
                f"## 1. Introduction\n\n"
                f"### 1.1 Scope\n"
                f"This document specifies the requirements and specifications for {project.name}. "
                f"It defines the system architecture, functional requirements, and technical standards to be followed.\n\n"
                f"### 1.2 Purpose\n"
                f"The purpose of this document is to provide a comprehensive technical specification that serves as:\n"
                f"- A reference for system design and implementation\n"
                f"- A basis for testing and validation\n"
                f"- A guide for maintenance and future enhancements\n\n"
                f"### 1.3 Document Control\n"
                f"- **Project**: {project.name}\n"
                f"- **Organization**: {project.organization.name}\n"
                f"- **Author**: @{request.user.username}\n"
                f"- **Version**: 1.0\n"
                f"- **Status**: DRAFT\n"
                f"- **Date**: {datetime.datetime.now().strftime('%Y-%m-%d')}\n\n"
                f"## 2. References\n\n"
                f"### 2.1 Applicable Documents\n"
                f"- IEEE Std 830-1998: IEEE Recommended Practice for Software Requirements Specifications\n"
                f"- IEEE Std 1012-2016: IEEE Standard for System and Software Verification and Validation\n"
                f"- Project-specific requirements and design documents\n\n"
                f"### 2.2 Definitions and Acronyms\n"
                f"- **API**: Application Programming Interface\n"
                f"- **UI**: User Interface\n"
                f"- **UX**: User Experience\n"
                f"- **SRS**: Software Requirements Specification\n\n"
                f"## 3. System Overview\n\n"
                f"### 3.1 System Description\n"
                f"{project.description or 'System description to be provided.'}\n\n"
                f"### 3.2 System Architecture\n"
                f"The system shall follow a modular architecture with the following components:\n"
                f"- Frontend user interface\n"
                f"- Backend API services\n"
                f"- Database layer\n"
                f"- Integration interfaces\n\n"
                f"## 4. Functional Requirements\n\n"
                f"### 4.1 User Interface Requirements\n"
                f"- **REQ-UI-001**: The system shall provide a responsive web interface\n"
                f"- **REQ-UI-002**: The system shall support role-based access control\n"
                f"- **REQ-UI-003**: The system shall provide real-time feedback for user actions\n\n"
                f"### 4.2 Business Logic Requirements\n"
                f"- **REQ-BL-001**: The system shall validate all user inputs before processing\n"
                f"- **REQ-BL-002**: The system shall maintain data consistency across all modules\n"
                f"- **REQ-BL-003**: The system shall support audit logging for all critical operations\n\n"
                f"### 4.3 Data Management Requirements\n"
                f"- **REQ-DM-001**: The system shall ensure data persistence through ACID-compliant transactions\n"
                f"- **REQ-DM-002**: The system shall implement data backup and recovery mechanisms\n"
                f"- **REQ-DM-003**: The system shall support data export in multiple formats\n\n"
                f"## 5. Non-Functional Requirements\n\n"
                f"### 5.1 Performance Requirements\n"
                f"- **REQ-PERF-001**: The system shall respond to user requests within 2 seconds under normal load\n"
                f"- **REQ-PERF-002**: The system shall support 1000 concurrent users\n"
                f"- **REQ-PERF-003**: The system shall maintain 99.9% uptime availability\n\n"
                f"### 5.2 Security Requirements\n"
                f"- **REQ-SEC-001**: The system shall authenticate all users via secure protocols\n"
                f"- **REQ-SEC-002**: The system shall encrypt sensitive data at rest and in transit\n"
                f"- **REQ-SEC-003**: The system shall implement role-based access control\n"
                f"- **REQ-SEC-004**: The system shall log all security-relevant events\n\n"
                f"### 5.3 Reliability Requirements\n"
                f"- **REQ-REL-001**: The system shall have mean time between failures (MTBF) of 720 hours\n"
                f"- **REQ-REL-002**: The system shall recover from failures within 5 minutes\n"
                f"- **REQ-REL-003**: The system shall implement graceful degradation under load\n\n"
                f"## 6. Interface Requirements\n\n"
                f"### 6.1 User Interfaces\n"
                f"The system shall provide web-based user interfaces accessible via modern browsers.\n\n"
                f"### 6.2 External Interfaces\n"
                f"- **API Interface**: RESTful API for third-party integrations\n"
                f"- **Authentication Interface**: OAuth 2.0 / SAML integration\n"
                f"- **Notification Interface**: Email and push notification services\n\n"
                f"## 7. Design Constraints\n\n"
                f"### 7.1 Technical Constraints\n"
                f"- The system shall be built using approved technology stack\n"
                f"- The system shall comply with organizational security policies\n"
                f"- The system shall support scalability requirements\n\n"
                f"### 7.2 Regulatory Constraints\n"
                f"- The system shall comply with data protection regulations\n"
                f"- The system shall maintain audit trails for compliance\n\n"
                f"## 8. Verification and Validation\n\n"
                f"### 8.1 Testing Strategy\n"
                f"- Unit testing for all components\n"
                f"- Integration testing for module interactions\n"
                f"- System testing for end-to-end functionality\n"
                f"- User acceptance testing for stakeholder validation\n\n"
                f"### 8.2 Acceptance Criteria\n"
                f"- All functional requirements must be implemented and tested\n"
                f"- All non-functional requirements must be met\n"
                f"- Security audit must be completed successfully\n"
                f"- Performance benchmarks must be achieved\n\n"
                f"## 9. Appendices\n\n"
                f"### Appendix A: Stakeholder Analysis\n"
                f"| Stakeholder | Role | Interest | Influence |\n"
                f"| --- | --- | --- | --- |\n"
            )
            
            if stakeholders.exists():
                for s in stakeholders:
                    content += f"| {s.name} | {s.title} | {s.interest} | {s.power} |\n"
            else:
                content += "| [Stakeholder Name] | [Role] | [Interest Level] | [Influence Level] |\n"
            
            content += "\n\n### Appendix B: Requirements Traceability\n"
            content += "| Requirement ID | Description | Priority | Status |\n"
            content += "| --- | --- | --- | --- |\n"
            
            if requirements.exists():
                for r in requirements:
                    content += f"| {r.req_id} | {r.title} | {r.priority} | {r.status} |\n"
            else:
                content += "| REQ-XXX | [Requirement Description] | [Priority] | [Status] |\n"
            
            content += "\n\n### Appendix C: Revision History\n"
            content += "| Version | Date | Author | Description |\n"
            content += "| --- | --- | --- | --- |\n"
            content += "| 1.0 | " + datetime.datetime.now().strftime('%Y-%m-%d') + f" | {request.user.username} | Initial draft |\n"
        else:
            title = f"{doc_type} - {project.name} - Version 1.0"
            content = f"# Business Document: {doc_type} - {project.name}\n\n"
            content += "## 1. Document Control & Scope\n"
            content += f"- **Project Scope**: {project.name}\n"
            content += f"- **Workspace Organisation**: {project.organization.name}\n"
            content += f"- **Document Schema Type**: {doc_type}\n"
            content += f"- **Author**: @{request.user.username}\n"
            content += "- **Status**: DRAFT\n"
            content += f"- **Scope Description**: {project.description or 'No scope definition provided.'}\n\n"

            content += "## 2. Key Stakeholder Registry\n"
            if not stakeholders.exists():
                content += "*No stakeholders registered in this project.*\n\n"
            else:
                content += "| Name | Title | Department | Power / Interest |\n"
                content += "| --- | --- | --- | --- |\n"
                for s in stakeholders:
                    content += f"| {s.name} | {s.title} | {s.department or 'N/A'} | {s.power} / {s.interest} |\n"
                content += "\n"

            content += "## 3. Business & Technical Specifications Catalog\n"
            if not requirements.exists():
                content += "*No specifications recorded in project backlog.*\n\n"
            else:
                content += "| ID | Title | Priority | Type | Status |\n"
                content += "| --- | --- | --- | --- | --- |\n"
                for r in requirements:
                    content += f"| {r.req_id} | {r.title} | {r.priority} | {r.req_type} | {r.status} |\n"
                content += "\n"

        if doc_type == "FRD":
            content += "## 4. Agile Backlog & User Stories Traceability\n"
            from stories.models import UserStory
            stories = UserStory.objects.filter(requirement__project=project)
            if not stories.exists():
                content += "*No user story mappings found in backlog.*\n\n"
            else:
                content += "| Story ID | Title | Estimation | Status | Traced Requirement |\n"
                content += "| --- | --- | --- | --- | --- |\n"
                for s in stories:
                    content += f"| {s.story_id} | {s.title} | {s.points} pts | {s.status} | {s.requirement.req_id} |\n"
                content += "\n"

        # Embed Project Diagrams and Models
        project_diagrams = project.diagrams.all()
        if project_diagrams.exists():
            content += "## 5. Visual Analysis Models & Diagram Specifications\n"
            for d in project_diagrams:
                # For FRD, prioritize sequence/activity/ERD diagrams; for BRD, include all process/usecase
                if doc_type == "FRD" and d.diagram_type.upper() not in ["SEQUENCE", "ERD", "BPMN 2.0", "USE CASE"]:
                    continue
                content += f"### {d.name} ({d.diagram_type})\n"
                if d.description:
                    content += f"**Description**: {d.description}\n\n"
                if d.documentation:
                    content += f"**Model Documentation Details**:\n{d.documentation}\n\n"
                
                try:
                    from diagrams.exporters import export_to_mermaid
                    mermaid_code = export_to_mermaid(d.diagram_type, d.canvas_json)
                    content += "```mermaid\n"
                    content += mermaid_code + "\n"
                    content += "```\n\n"
                except Exception as e:
                    content += f"*Unable to export visual graphic: {str(e)}*\n\n"

        title = f"{doc_type} - {project.name} - Version 1.0"

        data = {
            "project": str(project.id),
            "doc_type": doc_type,
            "title": title,
            "version": "1.0",
            "status": "DRAFT",
            "content": content,
        }

        return api_success(data=data, message="Document synthesized successfully.")

    @action(detail=False, methods=["post"], url_path="generate-with-ai")
    def generate_with_ai(self, request):
        """
        Uses AI to generate a document drafted based on the user's prompt and project context.
        """
        project_id = request.data.get("project")
        doc_type = request.data.get("doc_type", "BRD")
        user_prompt = request.data.get("prompt", "Generate standard document.")
        
        if not project_id:
            return api_error(message="Project ID is required.")

        try:
            project = Project.objects.get(
                id=project_id,
                organization_id=request.user.organization_id
            )
        except Project.DoesNotExist:
            return api_error(message="Project not found in your organization.")

        stakeholders = project.stakeholders.all()
        requirements = project.requirements.all()
        
        from stories.models import UserStory
        user_stories = UserStory.objects.filter(requirement__project=project)
        
        # Build context dictionary
        context = {
            "project_name": project.name,
            "project_description": project.description,
            "stakeholders": {s.name: f"{s.title} ({s.department})" for s in stakeholders} if stakeholders.exists() else {},
            "requirements": [{"title": r.title, "description": r.description} for r in requirements] if requirements.exists() else [],
            "user_stories": [{"title": s.title, "description": f"As a {s.role}, I want to {s.action} so that {s.benefit}"} for s in user_stories] if user_stories.exists() else []
        }
        
        try:
            import asyncio
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            from srs.ai_integration import ai_service
            
            enhanced_content = loop.run_until_complete(
                ai_service.generate_full_document(
                    doc_type=doc_type,
                    user_prompt=user_prompt,
                    context=context,
                    organization=request.user.organization,
                    user=request.user
                )
            )
            loop.close()
        except Exception as e:
            return api_error(message=f"AI document generation failed: {str(e)}")
            
        title = f"{doc_type} - {project.name} (AI Generated) - Version 1.0"
        
        data = {
            "project": str(project.id),
            "doc_type": doc_type,
            "title": title,
            "version": "1.0",
            "status": "DRAFT",
            "content": enhanced_content,
        }
        
        return api_success(data=data, message="AI Document drafted successfully.")

    @action(detail=True, methods=["post"], url_path="submit-for-review")
    def submit_for_review(self, request, pk=None):
        """
        Transitions document status to REVIEW.
        """
        doc = self.get_object()
        doc.status = "REVIEW"
        doc.save()
        
        DocumentApprovalHistory.objects.create(
            document=doc,
            user=request.user,
            action="SUBMIT_REVIEW",
            comment="Submitted for review & approvals cataloging.",
            version=doc.version
        )
        
        # Ensure a version snapshot exists representing this version state
        DocumentVersion.objects.get_or_create(
            document=doc,
            version=doc.version,
            defaults={"content": doc.content, "created_by": request.user}
        )
        
        serializer = self.get_serializer(doc)
        return api_success(data=serializer.data, message="Document submitted for review successfully.")

    @action(detail=True, methods=["post"], url_path="approve")
    def approve(self, request, pk=None):
        """
        Transitions document status to APPROVED.
        """
        doc = self.get_object()
        if request.user.role not in ["ADMIN", "PRODUCT_OWNER", "PROJECT_MANAGER"]:
            return api_error(message="Only Product Owners, Project Managers, and Admins can approve documents.")
            
        doc.status = "APPROVED"
        doc.save()
        
        DocumentApprovalHistory.objects.create(
            document=doc,
            user=request.user,
            action="APPROVE",
            comment="Approved business specification draft.",
            version=doc.version
        )
        
        serializer = self.get_serializer(doc)
        return api_success(data=serializer.data, message="Document approved successfully.")

    @action(detail=True, methods=["post"], url_path="request-revision")
    def request_revision(self, request, pk=None):
        """
        Transitions document status back to DRAFT with rejection review reasons.
        """
        doc = self.get_object()
        if request.user.role not in ["ADMIN", "PRODUCT_OWNER", "PROJECT_MANAGER"]:
            return api_error(message="Only Product Owners, Project Managers, and Admins can request document revisions.")
            
        comment = request.data.get("comment", "").strip()
        if not comment:
            return api_error(message="A revision comment/rejection reason is required.")
            
        doc.status = "DRAFT"
        doc.save()
        
        DocumentApprovalHistory.objects.create(
            document=doc,
            user=request.user,
            action="REQUEST_REVISIONS",
            comment=comment,
            version=doc.version
        )
        
        serializer = self.get_serializer(doc)
        return api_success(data=serializer.data, message="Document revisions requested successfully.")

    @action(detail=True, methods=["post"], url_path="sign-off")
    def sign_off(self, request, pk=None):
        """
        Transitions document status to SIGNED_OFF.
        Restricted to Admins, Product Owners, and Project Managers.
        """
        doc = self.get_object()

        if request.user.role not in ["ADMIN", "PRODUCT_OWNER", "PROJECT_MANAGER"]:
            return api_error(message="Only Product Owners, Project Managers, and Admins can sign off documents.")

        from django.utils import timezone
        doc.status = "SIGNED_OFF"
        doc.signed_off_by = request.user
        doc.signed_off_at = timezone.now()
        doc.save()

        DocumentApprovalHistory.objects.create(
            document=doc,
            user=request.user,
            action="SIGN_OFF",
            comment="Officially signed off and compliance-approved.",
            version=doc.version
        )

        serializer = self.get_serializer(doc)
        return api_success(data=serializer.data, message="Document signed off successfully.")

    @action(detail=True, methods=["get"], url_path="history")
    def history(self, request, pk=None):
        """
        Retrieves approval audit logs.
        """
        doc = self.get_object()
        histories = doc.approval_histories.all()
        from .serializers import DocumentApprovalHistorySerializer
        serializer = DocumentApprovalHistorySerializer(histories, many=True)
        return api_success(data=serializer.data, message="Approval history retrieved.")

    @action(detail=True, methods=["get"], url_path="versions")
    def versions(self, request, pk=None):
        """
        Retrieves historical snapshots.
        """
        doc = self.get_object()
        versions = doc.versions.all()
        from .serializers import DocumentVersionSerializer
        serializer = DocumentVersionSerializer(versions, many=True)
        return api_success(data=serializer.data, message="Historical versions retrieved.")

    @action(detail=True, methods=["post"], url_path="ai-enhance")
    def ai_enhance(self, request, pk=None):
        """
        Enhances document content using AI.
        Supports section expansion, refinement, and IEEE formatting.
        """
        if not AI_AVAILABLE:
            return api_error(message="AI service is not available. Please configure AI integration.")
        
        doc = self.get_object()
        enhancement_type = request.data.get("enhancement_type", "refine")  # expand, refine, format
        section_content = request.data.get("content", doc.content)
        
        try:
            # Build context for AI enhancement
            context = {
                "doc_type": doc.doc_type,
                "project": doc.project.name,
                "title": doc.title,
                "version": doc.version,
            }
            
            # For IEEE documents, add specific context
            if doc.doc_type == "IEEE":
                context["format_standard"] = "IEEE 830"
                context["document_structure"] = "technical specification"
            
            # Call AI service for enhancement (synchronous for now)
            import asyncio
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            enhanced_content = loop.run_until_complete(
                ai_service.enhance_section(
                    section_content=section_content,
                    enhancement_type=enhancement_type,
                    context=context
                )
            )
            
            loop.close()
            
            # Update document with enhanced content
            doc.content = enhanced_content
            doc.save()
            
            # Create version snapshot
            DocumentVersion.objects.create(
                document=doc,
                version=doc.version,
                content=enhanced_content,
                created_by=request.user
            )
            
            serializer = self.get_serializer(doc)
            return api_success(data=serializer.data, message="Document enhanced successfully using AI.")
            
        except Exception as e:
            return api_error(message=f"AI enhancement failed: {str(e)}")

    def list(self, request, *args, **kwargs):
        queryset = self.filter_queryset(self.get_queryset())
        serializer = self.get_serializer(queryset, many=True)
        return api_success(data=serializer.data, message="Documents retrieved successfully.")

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        serializer = self.get_serializer(instance)
        return api_success(data=serializer.data, message="Document details retrieved.")

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return api_success(
            data=serializer.data,
            message="Document created successfully.",
            status_code=status.HTTP_201_CREATED
        )

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop("partial", False)
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        self.perform_update(serializer)
        return api_success(data=serializer.data, message="Document updated successfully.")

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        self.perform_destroy(instance)
        return api_success(message="Document deleted successfully.", status_code=status.HTTP_200_OK)

    @action(detail=True, methods=["get"], url_path="export-pdf")
    def export_pdf(self, request, pk=None):
        document = self.get_object()
        pdf_io = io.BytesIO()
        self._generate_pdf_reportlab(document, pdf_io)
        pdf_io.seek(0)
        
        response = FileResponse(pdf_io, content_type='application/pdf')
        safe_title = "".join(x for x in document.title if x.isalnum() or x in "._- ")
        response['Content-Disposition'] = f'attachment; filename="{safe_title}.pdf"'
        return response

    @action(detail=True, methods=["get"], url_path="export-word")
    def export_word(self, request, pk=None):
        document = self.get_object()
        doc = self._markdown_to_docx(document.content)
        word_io = io.BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        
        response = FileResponse(word_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        safe_title = "".join(x for x in document.title if x.isalnum() or x in "._- ")
        response['Content-Disposition'] = f'attachment; filename="{safe_title}.docx"'
        return response

    @action(detail=True, methods=["get"], url_path="export-markdown")
    def export_markdown(self, request, pk=None):
        document = self.get_object()
        markdown_io = io.BytesIO(document.content.encode('utf-8'))
        markdown_io.seek(0)
        
        response = FileResponse(markdown_io, content_type='text/markdown')
        safe_title = "".join(x for x in document.title if x.isalnum() or x in "._- ")
        response['Content-Disposition'] = f'attachment; filename="{safe_title}.md"'
        return response

    @action(detail=False, methods=["post"], url_path="import-markdown")
    def import_markdown(self, request):
        """
        Import markdown content and create a new document.
        """
        import tempfile
        import os
        
        if 'file' not in request.FILES:
            return api_error(message="No file provided.")
        
        file = request.FILES['file']
        project_id = request.data.get("project")
        doc_type = request.data.get("doc_type", "BRD")
        title = request.data.get("title", file.name.replace('.md', ''))
        
        if not project_id:
            return api_error(message="Project ID is required.")
        
        try:
            project = Project.objects.get(
                id=project_id,
                organization_id=request.user.organization_id
            )
        except Project.DoesNotExist:
            return api_error(message="Project not found in your organization.")
        
        try:
            # Read markdown content
            content = file.read().decode('utf-8')
            
            # Create document
            document = BusinessDocument.objects.create(
                project=project,
                doc_type=doc_type,
                title=title,
                version="1.0",
                status="DRAFT",
                content=content,
                created_by=request.user
            )
            
            # Create version snapshot
            DocumentVersion.objects.create(
                document=document,
                version=document.version,
                content=content,
                created_by=request.user
            )
            
            serializer = self.get_serializer(document)
            return api_success(data=serializer.data, message="Markdown document imported successfully.")
            
        except Exception as e:
            return api_error(message=f"Failed to import markdown file: {str(e)}")

    @action(detail=False, methods=["post"], url_path="import-docx")
    def import_docx(self, request):
        """
        Import DOCX content and create a new document.
        """
        if 'file' not in request.FILES:
            return api_error(message="No file provided.")
        
        file = request.FILES['file']
        project_id = request.data.get("project")
        doc_type = request.data.get("doc_type", "BRD")
        title = request.data.get("title", file.name.replace('.docx', ''))
        
        if not project_id:
            return api_error(message="Project ID is required.")
        
        try:
            project = Project.objects.get(
                id=project_id,
                organization_id=request.user.organization_id
            )
        except Project.DoesNotExist:
            return api_error(message="Project not found in your organization.")
        
        try:
            # Read DOCX content
            doc = Document(file)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n\n"
            
            # Convert basic formatting to markdown
            content = content.replace('\n\n\n', '\n\n')
            
            # Create document
            document = BusinessDocument.objects.create(
                project=project,
                doc_type=doc_type,
                title=title,
                version="1.0",
                status="DRAFT",
                content=content,
                created_by=request.user
            )
            
            # Create version snapshot
            DocumentVersion.objects.create(
                document=document,
                version=document.version,
                content=content,
                created_by=request.user
            )
            
            serializer = self.get_serializer(document)
            return api_success(data=serializer.data, message="DOCX document imported successfully.")
            
        except Exception as e:
            return api_error(message=f"Failed to import DOCX file: {str(e)}")

    def _parse_markdown_blocks(self, content):
        lines = content.split('\n')
        blocks = []
        current_block = None
        
        for line in lines:
            stripped = line.strip()
            
            # Code block
            if stripped.startswith('```'):
                if current_block and current_block['type'] == 'code':
                    blocks.append(current_block)
                    current_block = None
                else:
                    if current_block:
                        blocks.append(current_block)
                    current_block = {'type': 'code', 'lines': [line]}
                continue
                
            if current_block and current_block['type'] == 'code':
                current_block['lines'].append(line)
                continue
                
            # Empty line
            if not stripped:
                if current_block:
                    blocks.append(current_block)
                    current_block = None
                continue
                
            if stripped.startswith('# '):
                if current_block:
                    blocks.append(current_block)
                blocks.append({'type': 'h1', 'lines': [stripped]})
                current_block = None
                continue
            elif stripped.startswith('## '):
                if current_block:
                    blocks.append(current_block)
                blocks.append({'type': 'h2', 'lines': [stripped]})
                current_block = None
                continue
            elif stripped.startswith('### '):
                if current_block:
                    blocks.append(current_block)
                blocks.append({'type': 'h3', 'lines': [stripped]})
                current_block = None
                continue
                
            if stripped.startswith('- ') or stripped.startswith('* '):
                if current_block and current_block['type'] == 'list':
                    current_block['lines'].append(stripped)
                else:
                    if current_block:
                        blocks.append(current_block)
                    current_block = {'type': 'list', 'lines': [stripped]}
                continue
                
            if stripped.startswith('|'):
                if current_block and current_block['type'] == 'table':
                    current_block['lines'].append(stripped)
                else:
                    if current_block:
                        blocks.append(current_block)
                    current_block = {'type': 'table', 'lines': [stripped]}
                continue
                
            if current_block and current_block['type'] == 'paragraph':
                current_block['lines'].append(line)
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'paragraph', 'lines': [line]}
                
        if current_block:
            blocks.append(current_block)
            
        return blocks

    def _generate_pdf_reportlab(self, document, stream):
        doc = SimpleDocTemplate(stream, pagesize=A4, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=colors.HexColor('#1E3A8A'),
            spaceAfter=15
        )
        h2_style = ParagraphStyle(
            'DocH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#2563EB'),
            spaceBefore=15,
            spaceAfter=10
        )
        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor('#334155'),
            spaceAfter=8
        )
        bullet_style = ParagraphStyle(
            'DocBullet',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13.5,
            leftIndent=15,
            firstLineIndent=-10,
            textColor=colors.HexColor('#334155'),
            spaceAfter=4
        )
        
        blocks = self._parse_markdown_blocks(document.content)
        
        for block in blocks:
            if block['type'] == 'h1':
                story.append(Paragraph(block['lines'][0][2:], title_style))
                story.append(Spacer(1, 10))
            elif block['type'] == 'h2':
                story.append(Paragraph(block['lines'][0][3:], h2_style))
                story.append(Spacer(1, 8))
            elif block['type'] == 'h3':
                story.append(Paragraph(block['lines'][0][4:], styles['Heading3']))
                story.append(Spacer(1, 6))
            elif block['type'] == 'list':
                for line in block['lines']:
                    text = line[2:] if (line.startswith('- ') or line.startswith('* ')) else line
                    story.append(Paragraph(f"&bull; {text}", bullet_style))
                story.append(Spacer(1, 6))
            elif block['type'] == 'table':
                table_rows = []
                for line in block['lines']:
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells and all(c.startswith('-') for c in cells):
                        continue
                    table_rows.append(cells)
                if table_rows:
                    formatted_rows = []
                    for row_idx, r_data in enumerate(table_rows):
                        formatted_row = []
                        for cell_val in r_data:
                            c_style = ParagraphStyle('Cell', parent=body_style, fontSize=8, leading=10)
                            if row_idx == 0:
                                c_style = ParagraphStyle('HeaderCell', parent=body_style, fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)
                            formatted_row.append(Paragraph(cell_val, c_style))
                        formatted_rows.append(formatted_row)
                    
                    t = Table(formatted_rows)
                    t_style = [
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                        ('TOPPADDING', (0, 0), (-1, -1), 5),
                        ('LEFTPADDING', (0, 0), (-1, -1), 5),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ]
                    for r_idx in range(1, len(table_rows)):
                        if r_idx % 2 == 0:
                            t_style.append(('BACKGROUND', (0, r_idx), (-1, r_idx), colors.HexColor('#F8FAFC')))
                    t.setStyle(TableStyle(t_style))
                    story.append(t)
                    story.append(Spacer(1, 10))
            elif block['type'] == 'code':
                is_mermaid = any('mermaid' in line for line in block['lines'])
                if is_mermaid:
                    nodes = {}
                    edges = []
                    for line in block['lines']:
                        stripped_line = line.strip()
                        if not stripped_line or 'mermaid' in stripped_line or stripped_line.startswith('```') or stripped_line.startswith('graph') or stripped_line.startswith('subgraph') or stripped_line == 'end':
                            continue
                        
                        node_match = re.match(r'^(\w+)\["([^"]+)"\]', stripped_line) or re.match(r'^(\w+)\("([^"]+)"\)', stripped_line) or re.match(r'^(\w+)\[([^\]]+)\]', stripped_line)
                        if node_match:
                            nodes[node_match.group(1)] = node_match.group(2).strip()
                            continue
                            
                        edge_match = re.match(r'^(\w+)\s*-->\s*\|([^|]+)\|\s*(\w+)', stripped_line)
                        if edge_match:
                            edges.append((edge_match.group(1), edge_match.group(3), edge_match.group(2).strip()))
                            continue
                            
                        edge_match_no_label = re.match(r'^(\w+)\s*-->\s*(\w+)', stripped_line)
                        if edge_match_no_label:
                            edges.append((edge_match_no_label.group(1), edge_match_no_label.group(2), ''))
                            
                    story.append(Paragraph("<b>Visual Model Specifications</b>", styles['Heading3']))
                    story.append(Spacer(1, 6))
                    
                    if nodes:
                        story.append(Paragraph("<b>Model Elements & Actors</b>", body_style))
                        el_data = [["ID", "Element Name"]]
                        for k, v in nodes.items():
                            el_data.append([k, v])
                        
                        formatted_el = []
                        for row_idx, r_data in enumerate(el_data):
                            formatted_row = []
                            for cell_val in r_data:
                                c_style = ParagraphStyle('Cell', parent=body_style, fontSize=8, leading=10)
                                if row_idx == 0:
                                    c_style = ParagraphStyle('HeaderCell', parent=body_style, fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)
                                formatted_row.append(Paragraph(cell_val, c_style))
                            formatted_el.append(formatted_row)
                            
                        t = Table(formatted_el, colWidths=[100, 350])
                        t.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4B5563')),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                            ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 8))
                        
                    if edges:
                        story.append(Paragraph("<b>Flows & Interactions</b>", body_style))
                        flow_data = [["Source", "Action / Flow", "Target"]]
                        for src, tgt, label in edges:
                            src_lbl = nodes.get(src, src)
                            tgt_lbl = nodes.get(tgt, tgt)
                            lbl = label if label else "triggers"
                            flow_data.append([src_lbl, lbl, tgt_lbl])
                            
                        formatted_flow = []
                        for row_idx, r_data in enumerate(flow_data):
                            formatted_row = []
                            for cell_val in r_data:
                                c_style = ParagraphStyle('Cell', parent=body_style, fontSize=8, leading=10)
                                if row_idx == 0:
                                    c_style = ParagraphStyle('HeaderCell', parent=body_style, fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)
                                formatted_row.append(Paragraph(cell_val, c_style))
                            formatted_flow.append(formatted_row)
                            
                        t = Table(formatted_flow, colWidths=[175, 100, 175])
                        t.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4B5563')),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                            ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 10))
                else:
                    code_content = '\n'.join(block['lines'][1:-1])
                    code_style = ParagraphStyle(
                        'CodeStyle',
                        parent=styles['Normal'],
                        fontName='Courier',
                        fontSize=8,
                        leading=10,
                        textColor=colors.HexColor('#475569'),
                        leftIndent=10,
                        spaceAfter=8
                    )
                    story.append(Paragraph(code_content.replace('\n', '<br/>'), code_style))
                    story.append(Spacer(1, 8))
            elif block['type'] == 'paragraph':
                story.append(Paragraph('\n'.join(block['lines']), body_style))
                story.append(Spacer(1, 6))
                
        doc.build(story)

    def _markdown_to_docx(self, markdown_text):
        doc = Document()
        blocks = self._parse_markdown_blocks(markdown_text)
        
        for block in blocks:
            if block['type'] == 'h1':
                doc.add_heading(block['lines'][0][2:], level=1)
            elif block['type'] == 'h2':
                doc.add_heading(block['lines'][0][3:], level=2)
            elif block['type'] == 'h3':
                doc.add_heading(block['lines'][0][4:], level=3)
            elif block['type'] == 'list':
                for line in block['lines']:
                    text = line[2:] if (line.startswith('- ') or line.startswith('* ')) else line
                    doc.add_paragraph(text, style='List Bullet')
            elif block['type'] == 'table':
                table_rows = []
                for line in block['lines']:
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells and all(c.startswith('-') for c in cells):
                        continue
                    table_rows.append(cells)
                if table_rows:
                    num_cols = len(table_rows[0])
                    table = doc.add_table(rows=0, cols=num_cols)
                    table.style = 'Table Grid'
                    for r_data in table_rows:
                        row_cells = table.add_row().cells
                        for col_idx, cell_value in enumerate(r_data):
                            if col_idx < len(row_cells):
                                row_cells[col_idx].text = cell_value
            elif block['type'] == 'code':
                is_mermaid = any('mermaid' in line for line in block['lines'])
                if is_mermaid:
                    nodes = {}
                    edges = []
                    for line in block['lines']:
                        stripped_line = line.strip()
                        if not stripped_line or 'mermaid' in stripped_line or stripped_line.startswith('```') or stripped_line.startswith('graph') or stripped_line.startswith('subgraph') or stripped_line == 'end':
                            continue
                        
                        node_match = re.match(r'^(\w+)\["([^"]+)"\]', stripped_line) or re.match(r'^(\w+)\("([^"]+)"\)', stripped_line) or re.match(r'^(\w+)\[([^\]]+)\]', stripped_line)
                        if node_match:
                            nodes[node_match.group(1)] = node_match.group(2).strip()
                            continue
                            
                        edge_match = re.match(r'^(\w+)\s*-->\s*\|([^|]+)\|\s*(\w+)', stripped_line)
                        if edge_match:
                            edges.append((edge_match.group(1), edge_match.group(3), edge_match.group(2).strip()))
                            continue
                            
                        edge_match_no_label = re.match(r'^(\w+)\s*-->\s*(\w+)', stripped_line)
                        if edge_match_no_label:
                            edges.append((edge_match_no_label.group(1), edge_match_no_label.group(2), ''))
                            
                    doc.add_heading("Visual Model Specifications", level=3)
                    
                    if nodes:
                        p = doc.add_paragraph()
                        p.add_run("Model Elements & Actors").bold = True
                        table = doc.add_table(rows=1, cols=2)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'ID'
                        hdr_cells[1].text = 'Element Name'
                        for k, v in nodes.items():
                            row_cells = table.add_row().cells
                            row_cells[0].text = k
                            row_cells[1].text = v
                            
                    if edges:
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run("Flows & Interactions").bold = True
                        table = doc.add_table(rows=1, cols=3)
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Source'
                        hdr_cells[1].text = 'Action / Flow'
                        hdr_cells[2].text = 'Target'
                        for src, tgt, label in edges:
                            row_cells = table.add_row().cells
                            row_cells[0].text = nodes.get(src, src)
                            row_cells[1].text = label if label else "triggers"
                            row_cells[2].text = nodes.get(tgt, tgt)
                else:
                    code_content = '\n'.join(block['lines'][1:-1])
                    p = doc.add_paragraph()
                    run = p.add_run(code_content)
                    run.font.name = 'Courier New'
            elif block['type'] == 'paragraph':
                doc.add_paragraph('\n'.join(block['lines']))
                
        return doc
